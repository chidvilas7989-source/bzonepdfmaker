from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
import os
import uuid
import shutil
import time
import re
from PIL import Image
from google import genai
from google.genai import types
from dotenv import load_dotenv
import markdown
from xhtml2pdf import pisa
from pypdf import PdfReader, PdfWriter

load_dotenv()

# Bypass corrupted TensorFlow in user environment
os.environ["USE_TF"] = "0"
os.environ["USE_TORCH"] = "1"

# Initialize Gemini
api_key = os.environ.get("GEMINI_API_KEY", "")
client = None
if api_key and api_key != "your_gemini_api_key_here":
    client = genai.Client(api_key=api_key)
else:
    print("WARNING: Valid GEMINI_API_KEY not found in .env. Gemini structuring will fail.")

disable_local_ocr = os.environ.get("DISABLE_LOCAL_OCR", "0").lower() in ("1", "true", "yes")

HAS_LOCAL_OCR = False
processor = None
model = None
device = None

if not disable_local_ocr:
    try:
        from transformers import TrOCRProcessor, VisionEncoderDecoderModel
        import torch
        HAS_LOCAL_OCR = True
    except ImportError:
        print("TrOCR libraries (transformers/torch) not installed. Bypassing local OCR.")

if HAS_LOCAL_OCR:
    print("Loading TrOCR model (this may take a minute on first run)...")
    try:
        processor = TrOCRProcessor.from_pretrained('microsoft/trocr-base-handwritten')
        model = VisionEncoderDecoderModel.from_pretrained('microsoft/trocr-base-handwritten')
        device = "cuda" if torch.cuda.is_available() else "cpu"
        model.to(device)
        print("TrOCR model loaded successfully.")
    except Exception as e:
        print(f"Error loading TrOCR: {e}")
        HAS_LOCAL_OCR = False
else:
    print("Using Gemini for OCR.")

app = FastAPI(title="Bzone PDF Maker Backend", description="API for Handwriting to PDF Platform")

# Configure CORS for frontend access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup directories
UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Helper: HTML Styling for Pisa PDF conversion
def get_styled_html(html_content, document_title="Bzone PDF Maker"):
    return f"""
    <html>
    <head>
    <style>
        @page {{
            size: letter;
            margin: 2cm;
            margin-bottom: 2.5cm;
            @frame footer_frame {{
                -pdf-frame-content: footer_content;
                left: 2cm;
                width: 17.59cm;
                bottom: 1cm;
                height: 1cm;
            }}
        }}
        body {{
            font-family: 'Helvetica', 'Arial', sans-serif;
            color: #1e293b;
            line-height: 1.6;
            font-size: 10pt;
        }}
        h1 {{
            font-size: 22pt;
            color: #0f172a;
            margin-top: 0px;
            margin-bottom: 12px;
            padding-bottom: 8px;
            border-bottom: 2px solid #8b5cf6;
        }}
        h2 {{
            font-size: 15pt;
            color: #1e293b;
            margin-top: 20px;
            margin-bottom: 10px;
            padding-bottom: 4px;
            border-bottom: 1px solid #e2e8f0;
        }}
        h3 {{
            font-size: 12pt;
            color: #475569;
            margin-top: 16px;
            margin-bottom: 8px;
        }}
        p {{
            margin-bottom: 12px;
        }}
        ul, ol {{
            margin-bottom: 12px;
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 4px;
        }}
        blockquote {{
            margin: 16px 0;
            padding: 10px 20px;
            background-color: #f8fafc;
            border-left: 4px solid #8b5cf6;
            color: #475569;
            font-style: italic;
            border-radius: 4px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 16px 0;
        }}
        th, td {{
            padding: 8px 12px;
            text-align: left;
            border: 1px solid #e2e8f0;
        }}
        th {{
            background-color: #f1f5f9;
            color: #0f172a;
            font-weight: bold;
        }}
        code {{
            font-family: 'Courier New', Courier, monospace;
            background-color: #f1f5f9;
            padding: 2px 4px;
            border-radius: 4px;
            font-size: 9pt;
        }}
        pre {{
            font-family: 'Courier New', Courier, monospace;
            background-color: #f1f5f9;
            padding: 12px;
            border-radius: 8px;
            margin: 16px 0;
        }}
        #footer_content {{
            border-top: 0.5px solid #cbd5e1;
            padding-top: 4px;
        }}
    </style>
    </head>
    <body>
        {html_content}
        
        <div id="footer_content">
            <table style="width: 100%; border: none; margin: 0; padding: 0; font-family: 'Helvetica', 'Arial', sans-serif; font-size: 8pt; color: #64748b;">
                <tr>
                    <td style="border: none; padding: 0; text-align: left;">{document_title}</td>
                    <td style="border: none; padding: 0; text-align: right;">Page <pdf:pagenumber> of <pdf:pagecount></td>
                </tr>
            </table>
        </div>
    </body>
    </html>
    """

# Helper: Extract text from DOCX files using python-docx
def extract_text_from_docx(file_path):
    import docx
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading docx: {e}")
        return ""

# Helper: Convert .doc to .docx on Windows using win32com
def convert_doc_to_docx_win32(doc_path, docx_path):
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16) # 16 is wdFormatXMLDocument (.docx)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"Error converting .doc to .docx via win32com: {e}")
        return False

# Helper: Convert Markdown to DOCX file
def markdown_to_docx(markdown_text, docx_path):
    import docx
    from docx.shared import Pt, Inches
    
    doc = docx.Document()
    lines = markdown_text.split('\n')
    in_code_block = False
    code_content = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
            
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
            
        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text_docx(p, text)
            i += 1
            continue
            
        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            match = re.match(r'^\d+\.\s(.*)', line.strip())
            text = match.group(1).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_text_docx(p, text)
            i += 1
            continue
            
        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            parsed_rows = []
            for tl in table_lines:
                cols = [c.strip() for c in tl.split('|')[1:-1]]
                if all(re.match(r'^:?-+:?$', c) for c in cols if c):
                    continue
                parsed_rows.append(cols)
                
            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                num_rows = len(parsed_rows)
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                for r_idx, row_data in enumerate(parsed_rows):
                    row = table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < len(row.cells):
                            p = row.cells[c_idx].paragraphs[0]
                            if r_idx == 0:
                                add_formatted_text_docx(p, cell_value, bold_all=True)
                            else:
                                add_formatted_text_docx(p, cell_value)
            continue
            
        # Empty line
        if not line.strip():
            i += 1
            continue
            
        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_text_docx(p, line.strip())
        i += 1

    doc.save(docx_path)

# Helper: Format inline bold and italic markers in python-docx
def add_formatted_text_docx(paragraph, text, bold_all=False):
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        else:
            run.text = token
            if bold_all:
                run.bold = True

# Helper: Convert PDF back to DOCX by extracting text
def pdf_to_docx(pdf_path, docx_path):
    import docx
    try:
        reader = PdfReader(pdf_path)
        doc = docx.Document()
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                if page_num > 0:
                    doc.add_page_break()
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
        doc.save(docx_path)
    except Exception as e:
        print(f"Error converting pdf to docx: {e}")

# Serve routes
@app.get("/")
def serve_index():
    return FileResponse("index.html")

@app.get("/styles.css")
def serve_css():
    return FileResponse("styles.css")

@app.get("/app.js")
def serve_js():
    return FileResponse("app.js")

@app.get("/logo.png")
def serve_logo():
    return FileResponse("logo.png")

@app.get("/api/status")
def read_root():
    return {"status": "ok", "message": "Scribe Backend is running"}

@app.get("/api/download/{filename}")
async def download_pdf(filename: str):
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    media_type = 'application/octet-stream'
    if filename.endswith('.pdf'):
        media_type = 'application/pdf'
    elif filename.endswith('.docx'):
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif filename.endswith('.md'):
        media_type = 'text/markdown'
        
    return FileResponse(
        path=file_path, 
        filename=filename, 
        media_type=media_type
    )

# --- Endpoint: Export / Recompile edited markdown to PDF/DOCX ---
@app.post("/api/export")
async def export_document(
    markdown_text: str = Form(...),
    export_format: str = Form(...), # "pdf", "docx", "md"
    document_title: str = Form("Bzone Document")
):
    job_id = str(uuid.uuid4())
    if export_format == "pdf":
        pdf_filename = f"export_{job_id}.pdf"
        pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
        html_content = markdown.markdown(markdown_text, extensions=['tables'])
        html_doc = get_styled_html(html_content, document_title)
        with open(pdf_path, "w+b") as result_file:
            pisa.CreatePDF(html_doc, dest=result_file)
        return JSONResponse({"status": "success", "file_url": f"/api/download/{pdf_filename}"})
        
    elif export_format == "docx":
        docx_filename = f"export_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        markdown_to_docx(markdown_text, docx_path)
        return JSONResponse({"status": "success", "file_url": f"/api/download/{docx_filename}"})
        
    elif export_format == "md":
        md_filename = f"export_{job_id}.md"
        md_path = os.path.join(OUTPUT_DIR, md_filename)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        return JSONResponse({"status": "success", "file_url": f"/api/download/{md_filename}"})
        
    else:
        raise HTTPException(status_code=400, detail="Invalid export format")

# --- Endpoint: Handwriting OCR (Supports Images, PDFs, DOC/DOCX, TXT) ---
@app.post("/api/process")
async def process_document(
    file: UploadFile = File(...),
    doc_type: str = Form("general"),
    custom_prompt: str = Form(None)
):
    content_type = file.content_type or ""
    filename_lower = file.filename.lower()
    ext = os.path.splitext(filename_lower)[1]
    
    if not content_type:
        if ext in ('.jpg', '.jpeg', '.png', '.webp'):
            content_type = "image/jpeg"
        elif ext == '.pdf':
            content_type = "application/pdf"
        elif ext in ('.doc', '.docx'):
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext in ('.txt', '.text'):
            content_type = "text/plain"
            
    is_image = content_type.startswith("image/") or ext in ('.jpg', '.jpeg', '.png', '.webp')
    is_pdf = content_type == "application/pdf" or ext == '.pdf'
    is_docx = ext == '.docx' or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    is_doc = ext == '.doc' or content_type == "application/msword"
    is_txt = ext in ('.txt', '.text') or content_type.startswith("text/")

    if not (is_image or is_pdf or is_docx or is_doc or is_txt):
        raise HTTPException(status_code=400, detail="Invalid file type. Supported: JPG, PNG, PDF, DOC, DOCX, TXT.")

    job_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{job_id}{ext}")

    # Save uploaded file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        # Build prompt instructions based on doc_type and custom_prompt
        doc_formatting_instructions = ""
        if doc_type == "meeting":
            doc_formatting_instructions = (
                "Format the output as professional Meeting Minutes. Include sections for:\n"
                "- **Date & Time** (extract or guess from context, or leave a line placeholder)\n"
                "- **Attendees** (if any are listed)\n"
                "- **Overview/Objective**\n"
                "- **Key Discussion Points**\n"
                "- **Decisions Made**\n"
                "- **Action Items** (as a clear bulleted list with checkboxes `- [ ]` or indicators, and owners if mentioned)"
            )
        elif doc_type == "study":
            doc_formatting_instructions = (
                "Format the output as clear, organized Study/Lecture Notes. Ensure it includes:\n"
                "- A clear main topic header\n"
                "- Key terms or formulas highlighted in bold with definitions/explanations\n"
                "- Main concepts organized using logical nested bullet lists or sections\n"
                "- A final **Summary** or takeaways section"
            )
        elif doc_type == "todo":
            doc_formatting_instructions = (
                "Format the output as a neat To-Do / Task List. Use markdown task checkboxes (`- [ ]`) for items to be done. "
                "Group items by category if appropriate, and mark high-priority items with bold indicators."
            )
        elif doc_type == "receipt":
            doc_formatting_instructions = (
                "Format the output as a structured Expense Report/Invoice. Organize items, quantities, unit prices, "
                "and total costs into a clean Markdown table with columns. Include separate lines for transaction date, "
                "vendor, total taxes, and grand total at the bottom if present."
            )
        else:
            doc_formatting_instructions = (
                "Format the output as a clean, structured Markdown document. Use appropriate headings, bullet points, "
                "and typographic emphasis (bold/italic) to make it highly legible."
            )

        if custom_prompt and custom_prompt.strip():
            doc_formatting_instructions += f"\n\nCRITICAL USER REQUEST: Follow these additional formatting and processing instructions:\n{custom_prompt.strip()}"

        extracted_text = ""
        structured_markdown = ""

        # Process document based on file type
        if is_image:
            image = Image.open(file_path).convert("RGB")
            if HAS_LOCAL_OCR and processor and model:
                print(f"Extracting text from image {file.filename} using TrOCR...")
                pixel_values = processor(image, return_tensors="pt").pixel_values.to(device)
                generated_ids = model.generate(pixel_values)
                extracted_text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
                
                if client:
                    print("Structuring TrOCR text with Gemini...")
                    prompt = (
                        f"You are a document structuring assistant. Take the following raw OCR text extracted from a handwritten page and format it according to these rules:\n"
                        f"{doc_formatting_instructions}\n"
                        f"Fix obvious typos, but preserve the original semantic meaning.\n\n"
                        f"Raw OCR Text:\n{extracted_text}"
                    )
                    response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                    structured_markdown = response.text or ""
                else:
                    structured_markdown = f"# Extracted Text\n\n{extracted_text}"
            else:
                if not client:
                    raise HTTPException(status_code=400, detail="Local OCR is disabled and Gemini API client is not initialized.")
                print(f"Extracting and structuring text from image {file.filename} using Gemini Multimodal...")
                prompt = (
                    f"You are an expert handwriting transcription and formatting engine. Transcribe the handwriting in this image and format the output directly into a clean, structured Markdown document according to these guidelines:\n"
                    f"{doc_formatting_instructions}\n"
                    f"Fix any obvious handwriting/spelling typos, and use appropriate headers, bullet points, or tables to represent the structure of the document beautifully."
                )
                response = client.models.generate_content(model='gemini-2.5-flash', contents=[image, prompt])
                structured_markdown = response.text or "# Empty Document\n\nCould not extract any content."

        elif is_pdf:
            if client:
                print(f"Processing PDF {file.filename} with Gemini multimodal...")
                with open(file_path, "rb") as f:
                    pdf_bytes = f.read()
                prompt = (
                    f"You are an expert document transcription and formatting engine. Transcribe the contents of this PDF file and format the output directly into a clean, structured Markdown document according to these guidelines:\n"
                    f"{doc_formatting_instructions}\n"
                    f"Fix any typos, and use appropriate headers, bullet points, or tables to represent the structure of the document beautifully."
                )
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=[types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"), prompt]
                )
                structured_markdown = response.text or "# Empty Document\n\nCould not transcribe PDF."
            else:
                print(f"Reading PDF text locally from {file.filename} using pypdf...")
                reader = PdfReader(file_path)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        extracted_text += t + "\n"
                structured_markdown = f"# Extracted PDF Text\n\n{extracted_text}"

        elif is_docx:
            print(f"Extracting text from DOCX {file.filename}...")
            extracted_text = extract_text_from_docx(file_path)
            if client and extracted_text.strip():
                prompt = (
                    f"You are a document structuring assistant. Take the following extracted text from a Word document and format it according to these rules:\n"
                    f"{doc_formatting_instructions}\n"
                    f"Fix obvious typos, but preserve the original semantic meaning.\n\n"
                    f"Raw Text:\n{extracted_text}"
                )
                response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                structured_markdown = response.text or extracted_text
            else:
                structured_markdown = f"# Extracted DOCX Text\n\n{extracted_text}"

        elif is_doc:
            print(f"Converting DOC {file.filename} to DOCX...")
            temp_docx_path = file_path + ".docx"
            success = convert_doc_to_docx_win32(file_path, temp_docx_path)
            if success and os.path.exists(temp_docx_path):
                extracted_text = extract_text_from_docx(temp_docx_path)
                try:
                    os.remove(temp_docx_path)
                except:
                    pass
            else:
                raise HTTPException(
                    status_code=400, 
                    detail="Could not process .doc file. win32com conversion failed or MS Word is not installed on the server. Please save your file as modern Word Document (.docx) or Plain Text (.txt) and try again."
                )
                
            if client and extracted_text.strip():
                prompt = (
                    f"You are a document structuring assistant. Take the following extracted text from a Word document and format it according to these rules:\n"
                    f"{doc_formatting_instructions}\n"
                    f"Fix obvious typos, but preserve the original semantic meaning.\n\n"
                    f"Raw Text:\n{extracted_text}"
                )
                response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                structured_markdown = response.text or extracted_text
            else:
                structured_markdown = f"# Extracted DOC Text\n\n{extracted_text}"

        elif is_txt:
            print(f"Reading text from TXT {file.filename}...")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted_text = f.read()
            if client and extracted_text.strip():
                prompt = (
                    f"You are a document structuring assistant. Take the following raw text and format it according to these rules:\n"
                    f"{doc_formatting_instructions}\n"
                    f"Fix obvious typos, but preserve the original semantic meaning.\n\n"
                    f"Raw Text:\n{extracted_text}"
                )
                response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
                structured_markdown = response.text or extracted_text
            else:
                structured_markdown = f"# Text Content\n\n{extracted_text}"

        # 3. Generate outputs
        pdf_filename = f"{job_id}.pdf"
        pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
        html_content = markdown.markdown(structured_markdown, extensions=['tables'])
        html_doc = get_styled_html(html_content, "Bzone PDF Maker")
        with open(pdf_path, "w+b") as result_file:
            pisa.CreatePDF(html_doc, dest=result_file)

        docx_filename = f"{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        markdown_to_docx(structured_markdown, docx_path)

        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "markdown": structured_markdown,
            "pdf_url": f"/api/download/{pdf_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })

    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Endpoint: Merge PDFs ---
@app.post("/api/merge")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Please upload at least 2 PDF files to merge.")
    
    job_id = str(uuid.uuid4())
    output_filename = f"merged_{job_id}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    writer = PdfWriter()
    temp_files = []
    
    try:
        for file in files:
            if not file.filename.lower().endswith(".pdf"):
                raise HTTPException(status_code=400, detail="Only PDF files are supported for merging.")
            temp_file_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}.pdf")
            temp_files.append(temp_file_path)
            with open(temp_file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            reader = PdfReader(temp_file_path)
            for page in reader.pages:
                writer.add_page(page)
                
        with open(output_path, "wb") as out_file:
            writer.write(out_file)
            
        # Also convert merged PDF to docx format
        docx_filename = f"merged_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_to_docx(output_path, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "pdf_url": f"/api/download/{output_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass

# --- Endpoint: Split PDF ---
@app.post("/api/split")
async def split_pdf(
    file: UploadFile = File(...),
    pages_str: str = Form(...)  # e.g., "1, 2-5, 8"
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    
    job_id = str(uuid.uuid4())
    temp_file_path = os.path.join(UPLOAD_DIR, f"{job_id}.pdf")
    output_filename = f"split_{job_id}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        reader = PdfReader(temp_file_path)
        total_pages = len(reader.pages)
        writer = PdfWriter()
        
        selected_pages = []
        parts = pages_str.split(",")
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if "-" in part:
                start, end = part.split("-")
                start_idx = int(start.strip()) - 1
                end_idx = int(end.strip()) - 1
                start_idx = max(0, min(start_idx, total_pages - 1))
                end_idx = max(0, min(end_idx, total_pages - 1))
                for i in range(start_idx, end_idx + 1):
                    selected_pages.append(i)
            else:
                idx = int(part) - 1
                if 0 <= idx < total_pages:
                    selected_pages.append(idx)
                    
        if not selected_pages:
            raise HTTPException(status_code=400, detail="No valid pages specified for extraction.")
            
        for idx in selected_pages:
            writer.add_page(reader.pages[idx])
            
        with open(output_path, "wb") as out_file:
            writer.write(out_file)
            
        # Convert split PDF to docx format
        docx_filename = f"split_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_to_docx(output_path, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "pdf_url": f"/api/download/{output_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass

# --- Endpoint: Rotate PDF ---
@app.post("/api/rotate")
async def rotate_pdf(
    file: UploadFile = File(...),
    degrees: int = Form(...)  # 90, 180, 270
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    if degrees not in (90, 180, 270):
        raise HTTPException(status_code=400, detail="Degrees must be 90, 180, or 270.")
        
    job_id = str(uuid.uuid4())
    temp_file_path = os.path.join(UPLOAD_DIR, f"{job_id}.pdf")
    output_filename = f"rotated_{job_id}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        reader = PdfReader(temp_file_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            page.rotate(degrees)
            writer.add_page(page)
            
        with open(output_path, "wb") as out_file:
            writer.write(out_file)
            
        # Convert rotated PDF to docx
        docx_filename = f"rotated_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_to_docx(output_path, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "pdf_url": f"/api/download/{output_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass

# --- Endpoint: Protect PDF (Encrypt) ---
@app.post("/api/protect")
async def protect_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    if not password:
        raise HTTPException(status_code=400, detail="Password cannot be empty.")
        
    job_id = str(uuid.uuid4())
    temp_file_path = os.path.join(UPLOAD_DIR, f"{job_id}.pdf")
    output_filename = f"protected_{job_id}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        reader = PdfReader(temp_file_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
            
        writer.encrypt(password)
        
        with open(output_path, "wb") as out_file:
            writer.write(out_file)
            
        # Convert protected PDF to docx format
        docx_filename = f"protected_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_to_docx(output_path, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "pdf_url": f"/api/download/{output_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass

# --- Endpoint: Unlock PDF (Decrypt) ---
@app.post("/api/unlock")
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
        
    job_id = str(uuid.uuid4())
    temp_file_path = os.path.join(UPLOAD_DIR, f"{job_id}.pdf")
    output_filename = f"unlocked_{job_id}.pdf"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        reader = PdfReader(temp_file_path)
        if reader.is_encrypted:
            status = reader.decrypt(password)
            if not status:
                raise HTTPException(status_code=400, detail="Invalid password.")
                
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
            
        with open(output_path, "wb") as out_file:
            writer.write(out_file)
            
        # Convert unlocked PDF to docx format
        docx_filename = f"unlocked_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_to_docx(output_path, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "pdf_url": f"/api/download/{output_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to decrypt or read file. Ensure it is a valid encrypted PDF: {str(e)}")
    finally:
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass

# --- Endpoint: Summarize PDF ---
@app.post("/api/summarize")
async def summarize_pdf(
    file: UploadFile = File(...),
    custom_prompt: str = Form(None)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    if not client:
        raise HTTPException(status_code=400, detail="Gemini client is not initialized.")
        
    job_id = str(uuid.uuid4())
    temp_file_path = os.path.join(UPLOAD_DIR, f"{job_id}.pdf")
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        with open(temp_file_path, "rb") as f:
            pdf_bytes = f.read()
            
        prompt = "You are a document summarizing assistant. Summarize the attached PDF document into a clean, professional, and well-structured Markdown document. Use clear headings, key bullet points, and highlight important figures or dates."
        if custom_prompt and custom_prompt.strip():
            prompt += f"\n\nCRITICAL USER REQUEST: Follow these additional formatting and processing instructions:\n{custom_prompt.strip()}"
            
        print("Summarizing PDF with Gemini...")
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
                prompt
            ]
        )
        
        structured_markdown = response.text or "# Empty Summary\n\nCould not summarize this document."
        
        # Generate summary PDF
        pdf_filename = f"summary_{job_id}.pdf"
        pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
        html_content = markdown.markdown(structured_markdown, extensions=['tables'])
        html_doc = get_styled_html(html_content, "Bzone PDF Summarizer")
        with open(pdf_path, "w+b") as result_file:
            pisa.CreatePDF(html_doc, dest=result_file)
            
        # Generate summary DOCX
        docx_filename = f"summary_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        markdown_to_docx(structured_markdown, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "markdown": structured_markdown,
            "pdf_url": f"/api/download/{pdf_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass

# --- Endpoint: Translate PDF ---
@app.post("/api/translate")
async def translate_pdf(
    file: UploadFile = File(...),
    target_language: str = Form(...),
    custom_prompt: str = Form(None)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    if not client:
        raise HTTPException(status_code=400, detail="Gemini client is not initialized.")
        
    job_id = str(uuid.uuid4())
    temp_file_path = os.path.join(UPLOAD_DIR, f"{job_id}.pdf")
    
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        with open(temp_file_path, "rb") as f:
            pdf_bytes = f.read()
            
        prompt = f"You are a professional document translator. Translate the attached PDF document entirely into the following target language: {target_language}. Maintain the markdown structure and formatting. Do not output anything else but the translation."
        if custom_prompt and custom_prompt.strip():
            prompt += f"\n\nCRITICAL USER REQUEST: Follow these additional formatting and processing instructions:\n{custom_prompt.strip()}"
            
        print(f"Translating PDF to {target_language} with Gemini...")
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
                prompt
            ]
        )
        
        structured_markdown = response.text or "# Empty Translation\n\nTranslation failed."
        
        # Generate translated PDF
        pdf_filename = f"translated_{job_id}.pdf"
        pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
        html_content = markdown.markdown(structured_markdown, extensions=['tables'])
        html_doc = get_styled_html(html_content, f"Bzone PDF Translation ({target_language})")
        with open(pdf_path, "w+b") as result_file:
            pisa.CreatePDF(html_doc, dest=result_file)
            
        # Generate translated DOCX
        docx_filename = f"translated_{job_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        markdown_to_docx(structured_markdown, docx_path)
            
        return JSONResponse({
            "status": "success",
            "job_id": job_id,
            "markdown": structured_markdown,
            "pdf_url": f"/api/download/{pdf_filename}",
            "docx_url": f"/api/download/{docx_filename}"
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except:
                pass
